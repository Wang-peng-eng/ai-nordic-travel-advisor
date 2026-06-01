import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
JSON_PATH = ROOT / "output" / "test_results_three_tier.json"
DESKTOP = Path.home() / "Desktop"
OUTPUT = DESKTOP / "AI北欧定制游顾问_五组案例三档方案测试报告.docx"

AGENCY = {
    "name": "Nordic Horizon International Travel Services",
    "cn_name": "北境环旅国际旅游服务中心",
    "phone": "+86 21 6188 4572",
    "email": "visa.docs@nordichorizon.cn",
    "dept": "签证文书协调部",
}

TIER_LABELS = {"经济型": "经济优选", "舒适型": "舒适精选", "高端型": "高端臻享"}
TIER_COLORS = {
    "经济型": RGBColor(44, 92, 77),
    "舒适型": RGBColor(31, 95, 115),
    "高端型": RGBColor(180, 130, 50),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=Pt(9.5), color=None, alignment=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if alignment:
        paragraph.alignment = alignment
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = size
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, widths=None, header=False, col_count=None):
    if not rows:
        return
    actual_cols = col_count or len(rows[0])
    table = doc.add_table(rows=len(rows), cols=actual_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row[:actual_cols]):
            cell = table.cell(row_idx, col_idx)
            is_header = header and row_idx == 0
            set_cell_text(cell, value, bold=is_header, size=Pt(9))
            if is_header:
                set_cell_shading(cell, "1F5F73")
                set_cell_text(cell, value, bold=True, size=Pt(9), color=RGBColor(255, 255, 255))
            if widths and col_idx < len(widths):
                cell.width = widths[col_idx]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 1:
            run.font.color.rgb = RGBColor(23, 32, 42)
        elif level == 2:
            run.font.color.rgb = RGBColor(31, 95, 115)
        else:
            run.font.color.rgb = RGBColor(44, 92, 77)
    return paragraph


def add_body(doc, text, bold=False, size=Pt(10.5)):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = size
    return paragraph


def add_separator(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CFD8DC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_three_tier_comparison_table(doc, plans):
    """Build a comparison table showing all 3 tiers side by side."""
    tiers = ["经济型", "舒适型", "高端型"]
    plan_map = {p["tier"]: p for p in plans}

    # Header row
    header = ["对比维度"]
    for tier in tiers:
        plan = plan_map.get(tier, {})
        header.append(f'{TIER_LABELS[tier]}\n{plan.get("totalPrice", "")}')

    # Data rows
    dimension_keys = ["交通方式", "导游服务", "餐饮标准", "门票覆盖", "特色项目"]

    rows = [header]
    for dim in dimension_keys:
        row = [dim]
        for tier in tiers:
            plan = plan_map.get(tier, {})
            features = plan.get("featureRows", [])
            val = next((f[1] for f in features if f[0] == dim), "")
            row.append(val)
        rows.append(row)

    # Build table
    col_count = 4
    widths = [Cm(2.8), Cm(4.2), Cm(4.2), Cm(4.2)]
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                color = TIER_COLORS.get(tiers[col_idx - 1], RGBColor(100, 100, 100)) if col_idx > 0 else RGBColor(23, 32, 42)
                bg = "1F5F73" if col_idx == 0 else ("2C5C4D" if col_idx == 1 else ("1F5F73" if col_idx == 2 else "B48232"))
                set_cell_shading(cell, bg)
                set_cell_text(cell, value, bold=True, size=Pt(9.5), color=RGBColor(255, 255, 255),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else None)
            elif col_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                set_cell_text(cell, value, bold=True, size=Pt(9))
            else:
                set_cell_text(cell, value, size=Pt(9))
            if widths and col_idx < len(widths):
                cell.width = widths[col_idx]

    doc.add_paragraph()
    return table


def add_letter_block(doc, letter, title):
    """Add a letter section to the document."""
    add_heading(doc, title, 2)
    if not letter:
        add_body(doc, "(无)")
        return
    add_heading(doc, letter.get("heading", ""), 3)
    add_body(doc, letter.get("recipient", ""))
    add_body(doc, letter.get("subject", ""), bold=True)
    for para in letter.get("paragraphs", []):
        add_body(doc, para)
    letter_table = letter.get("table", [])
    if letter_table:
        add_table(doc, letter_table, widths=[Cm(4.2), Cm(11.5)])
    add_body(doc, letter.get("closing", ""))


def write_case(doc, item, index):
    """Write a single case with all three tiers."""
    input_data = item["input"]
    plans = item["plans"]
    checks = item["checks"]
    plan_map = {p["tier"]: p for p in plans}

    # New page for each case (except first)
    if index > 1:
        doc.add_section(WD_SECTION_START.NEW_PAGE)

    # Case header
    add_heading(doc, f"案例 {index}：{input_data['customerName']}", 1)
    add_separator(doc)

    # --- Customer requirements ---
    add_heading(doc, "一、客户需求概览", 2)
    comfort = plan_map.get("舒适型", plans[1])
    summary = comfort["summary"]

    req_rows = [
        ["客户姓名", input_data["customerName"]],
        ["出行人数", f'{input_data["travelers"]}人'],
        ["出发城市", input_data["departureCity"]],
        ["目的地", input_data["destinations"]],
        ["行程天数", f'{input_data["days"]}天'],
        ["出发日期", input_data["departureDate"]],
        ["人均预算", f'¥{input_data["budget"]:,}'],
        ["出行偏好", input_data["preferences"]],
        ["特殊需求", input_data["specialNeeds"]],
        ["识别目的地", summary["country"]],
        ["识别客户类型", summary["customerType"]],
    ]
    add_table(doc, req_rows, widths=[Cm(3.5), Cm(12.3)])

    # --- Three-tier comparison ---
    add_heading(doc, "二、三档方案对比", 2)
    add_body(doc, "以下三档方案基于同一客户需求自动生成，在酒店等级、交通方式、导游服务、餐饮标准和特色项目上逐级升级。")
    build_three_tier_comparison_table(doc, plans)

    # --- Cost comparison ---
    add_heading(doc, "三、三档费用对比", 2)
    cost_header = ["费用项", "经济优选", "舒适精选", "高端臻享"]
    cost_dimensions = ["住宿费用", "交通费用", "活动费用", "服务费", "签证辅助费", "方案总价（整团）", "人均参考价"]
    cost_rows = [cost_header]
    for dim in cost_dimensions:
        row = [dim]
        for tier in ["经济型", "舒适型", "高端型"]:
            p = plan_map.get(tier, {})
            s = p.get("summary", {})
            costs = s.get("costsRows", [])
            val = next((c[1] for c in costs if c[0] == dim), "—")
            row.append(val)
        cost_rows.append(row)

    cost_table = doc.add_table(rows=len(cost_rows), cols=4)
    cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cost_table.style = "Table Grid"
    for row_idx, row in enumerate(cost_rows):
        for col_idx, value in enumerate(row):
            cell = cost_table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, "1F5F73")
                set_cell_text(cell, value, bold=True, size=Pt(9), color=RGBColor(255, 255, 255))
            elif col_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                set_cell_text(cell, value, bold=True, size=Pt(9))
            elif row_idx >= len(cost_rows) - 2:
                set_cell_text(cell, value, bold=True, size=Pt(10), color=RGBColor(180, 50, 50) if row_idx == len(cost_rows) - 1 else None)
            else:
                set_cell_text(cell, value, size=Pt(9))
    doc.add_paragraph()

    # --- Itinerary for recommended tier (舒适型) ---
    add_heading(doc, f"四、推荐方案行程安排（{comfort['tierLabel']}）", 2)
    add_body(doc, f"总价：{comfort['totalPrice']}　|　{summary['transportMode']}　|　{summary['guideType']}",
             bold=True)
    itinerary_section = next((s for s in comfort["sections"] if s["id"] == "itinerary"), None)
    if itinerary_section:
        add_table(doc, itinerary_section["rows"], widths=[Cm(1.5), Cm(14.3)])

    # --- Hotel confirmation letter (舒适型) ---
    add_heading(doc, "五、酒店确认函（舒适精选）", 2)
    add_letter_block(doc, summary.get("hotelLetter"), "HOTEL RESERVATION CONFIRMATION LETTER")

    # --- Invitation letter (舒适型) ---
    invitation_title = summary.get("invitationTitle", "邀请函")
    add_heading(doc, f"六、{invitation_title}（舒适精选）", 2)
    add_letter_block(doc, summary.get("invitationLetter"), invitation_title)

    # --- Visa checklist ---
    visa_section = next((s for s in comfort["sections"] if s["id"] == "visa"), None)
    if visa_section:
        add_heading(doc, "七、签证材料清单", 2)
        add_table(doc, visa_section["rows"], widths=[Cm(3), Cm(12.8)])

    # --- Test validation ---
    add_heading(doc, "八、测试校验", 2)
    check_rows = [[k, "✅ 通过" if v else "❌ 未通过"] for k, v in checks.items()]
    # Translate check keys
    key_labels = {
        "dayCountMatches": "行程天数与输入一致",
        "denmarkNoIcelandLeak": "丹麦案例无冰岛数据泄漏",
        "finlandNoIcelandLeak": "芬兰案例无冰岛数据泄漏",
        "icelandNoFinlandLeak": "冰岛案例无芬兰数据泄漏",
        "businessInviteOnlyWhenBusiness": "商务类型正确生成商务邀请函",
    }
    check_rows_labeled = [[key_labels.get(k, k), v] for k, v in check_rows]
    add_table(doc, check_rows_labeled, widths=[Cm(8), Cm(4)])


def build_document():
    with open(JSON_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Styles
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    # --- Cover page ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI北欧定制游顾问")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(23, 32, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("五组案例 · 三档方案自动生成测试报告")
    sub_run.font.name = "Microsoft YaHei"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(31, 95, 115)

    doc.add_paragraph()
    meta_rows = [
        ["签发机构", f"{AGENCY['name']} / {AGENCY['cn_name']}"],
        ["生成日期", datetime.now().strftime("%Y年%m月%d日 %H:%M")],
        ["文档编号", datetime.now().strftime("NH-THREE-TIER-%Y%m%d-%H%M")],
        ["联系方式", f"{AGENCY['phone']} / {AGENCY['email']}"],
        ["测试样本", f"{len(data)} 组随机客户案例 × 3 档方案 = {len(data) * 3} 套完整文书"],
        ["备注说明", "本报告由 AI北欧定制游顾问 系统自动生成，用于展示三档方案生成能力。"],
    ]
    add_table(doc, meta_rows, widths=[Cm(3.5), Cm(12.3)])

    # Overview summary
    add_heading(doc, "测试结论总览", 1)
    all_pass = all(all(c["checks"].values()) for c in data)
    add_body(doc, "✅ 全部案例通过校验" if all_pass else "⚠️ 存在需复核项", bold=True, size=Pt(12))

    summary_text = (
        f"本次测试覆盖 {len(data)} 组不同客户画像（含不同目的地国家、不同出行偏好、不同预算区间），"
        "每组均自动生成经济优选、舒适精选、高端臻享三档完整方案。"
        "校验项包括：行程天数匹配、目的地知识库隔离（无跨国家数据泄漏）、客户类型与邀请函类型匹配。"
    )
    add_body(doc, summary_text)

    # Summary table
    all_checks_pass = all_pass
    add_body(doc, f"\n校验结果：{'全部通过' if all_checks_pass else '存在未通过项'}", bold=True)
    summary_header = ["#", "客户", "目的地", "客户类型", "预算", "经济型总价", "舒适型总价", "高端型总价", "校验"]
    summary_rows = [summary_header]
    for i, item in enumerate(data, 1):
        inp = item["input"]
        plan_map = {p["tier"]: p for p in item["plans"]}
        comfort = plan_map.get("舒适型", item["plans"][1])
        ok = all(item["checks"].values())
        summary_rows.append([
            str(i),
            inp["customerName"],
            comfort["summary"]["country"],
            comfort["summary"]["customerType"],
            f'¥{inp["budget"]:,}',
            plan_map.get("经济型", {}).get("totalPrice", "—"),
            plan_map.get("舒适型", {}).get("totalPrice", "—"),
            plan_map.get("高端型", {}).get("totalPrice", "—"),
            "✅" if ok else "❌",
        ])
    add_table(doc, summary_rows, widths=[Cm(0.8), Cm(2), Cm(2), Cm(2.2), Cm(2), Cm(2.2), Cm(2.2), Cm(2.2), Cm(1.2)],
              header=True, col_count=9)

    # --- Individual cases ---
    for i, item in enumerate(data, 1):
        write_case(doc, item, i)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"报告已生成: {path}")
