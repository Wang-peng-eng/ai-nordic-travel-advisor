import json
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "AI北欧定制游顾问_五组案例测试结果.docx"


def run_cases():
    completed = subprocess.run(
        ["node", "scripts/run_case_tests.mjs"],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return json.loads(completed.stdout)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, widths=None, header=False):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, value, bold=header and row_idx == 0)
            if header and row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            if widths:
                cell.width = widths[col_idx]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
      run.font.name = "Microsoft YaHei"
      run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
      run.font.color.rgb = RGBColor(31, 77, 120 if level > 1 else 181)
    return paragraph


def add_body(doc, text, bold=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return paragraph


def add_letter(doc, letter):
    add_body(doc, letter["heading"], bold=True)
    add_body(doc, letter["recipient"])
    add_body(doc, letter["subject"], bold=True)
    for paragraph in letter["paragraphs"]:
        add_body(doc, paragraph)
    add_table(doc, letter["table"], widths=[Cm(4.2), Cm(11.5)])
    add_body(doc, letter["closing"])


def build_document(results):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI北欧定制游顾问 - 五组案例测试报告")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("目的地知识库、酒店匹配、客户类型识别与动态行程生成验证")
    subtitle_run.font.name = "Microsoft YaHei"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc_number = datetime.now().strftime("NH-TEST-%Y%m%d-%H%M")
    add_table(
        doc,
        [
            ["文档编号", doc_number],
            ["生成日期", generated_at],
            ["联系方式", "+86 21 6188 4572 / visa.docs@nordichorizon.cn"],
            ["备注说明", "本报告由当前项目生成逻辑直接执行五组测试案例后生成，用于内部验收。"],
        ],
        widths=[Cm(4), Cm(12)],
    )

    add_heading(doc, "测试结论", 1)
    all_checks_pass = all(all(item["checks"].values()) for item in results)
    add_body(doc, "通过" if all_checks_pass else "存在需复核项", bold=True)
    add_body(doc, "五组案例均按输入目的地读取对应知识库，按行程天数生成 Day 数量，按预算匹配酒店等级，并按客户类型生成不同邀请/说明函。")

    summary_rows = [["客户", "识别目的地", "客户类型", "Day 数", "酒店匹配", "邀请函类型", "验证"]]
    for item in results:
        input_data = item["input"]
        summary = item["summary"]
        checks_ok = all(item["checks"].values())
        hotel = "；".join(row[1] for row in summary["hotelRows"])
        summary_rows.append([
            input_data["customerName"],
            summary["country"],
            summary["customerType"],
            f'{summary["dayCount"]} / {input_data["days"]}',
            hotel,
            summary["invitationTitle"],
            "通过" if checks_ok else "需复核",
        ])
    add_table(doc, summary_rows, widths=[Cm(2), Cm(2), Cm(2.4), Cm(1.8), Cm(5), Cm(2.5), Cm(1.6)], header=True)

    for index, item in enumerate(results, start=1):
        input_data = item["input"]
        summary = item["summary"]
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_heading(doc, f"案例 {index}：{input_data['customerName']}", 1)

        add_heading(doc, "客户需求概览", 2)
        add_table(
            doc,
            [
                ["客户姓名", input_data["customerName"]],
                ["出行人数", f'{input_data["travelers"]}人'],
                ["出发城市", input_data["departureCity"]],
                ["目的地", input_data["destinations"]],
                ["出发日期", input_data["departureDate"]],
                ["行程天数", f'{input_data["days"]}天'],
                ["人均预算", f'¥{input_data["budget"]:,}'],
                ["出行偏好", input_data["preferences"]],
                ["特殊需求", input_data["specialNeeds"]],
                ["识别结果", f'{summary["country"]} / {summary["customerType"]} / {summary["hotelTier"]}'],
            ],
            widths=[Cm(3.2), Cm(12.5)],
        )

        add_heading(doc, "行程安排", 2)
        itinerary = next(section for section in item["sections"] if section["id"] == "itinerary")
        add_table(doc, itinerary["rows"], widths=[Cm(2), Cm(13.8)])

        add_heading(doc, "费用说明", 2)
        costs = next(section for section in item["sections"] if section["id"] == "costs")
        add_table(doc, costs["rows"], widths=[Cm(4), Cm(11.8)])
        for clause in costs.get("clauses", []):
            add_body(doc, clause)

        add_heading(doc, "签证材料清单", 2)
        visa = next(section for section in item["sections"] if section["id"] == "visa")
        add_table(doc, visa["rows"], widths=[Cm(4), Cm(11.8)])

        add_heading(doc, "酒店确认函", 2)
        add_letter(doc, summary["hotelLetter"])

        add_heading(doc, summary["invitationTitle"], 2)
        add_letter(doc, summary["invitationLetter"])

        add_heading(doc, "测试校验", 2)
        add_table(
            doc,
            [[key, "通过" if value else "未通过"] for key, value in item["checks"].items()],
            widths=[Cm(8), Cm(4)],
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    results = run_cases()
    path = build_document(results)
    print(path)
