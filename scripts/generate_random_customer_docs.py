from __future__ import annotations

import random
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


AGENCY = {
    "name": "Nordic Horizon International Travel Services",
    "cn_name": "北境环旅国际旅游服务中心",
    "phone": "+86 21 6188 4572",
    "email": "visa.docs@nordichorizon.cn",
    "dept": "签证文书协调部",
}

DESTINATIONS = {
    "冰岛": {
        "route": "雷克雅未克、黄金圈、南岸、杰古沙龙冰河湖",
        "entry": "雷克雅未克",
        "regions": ["雷克雅未克", "黄金圈", "南岸", "瓦特纳冰川", "斯奈山半岛"],
        "spots": {
            "自然风光": ["黄金瀑布", "辛格维利尔国家公园", "盖歇尔间歇泉", "塞里雅兰瀑布", "杰古沙龙冰河湖"],
            "摄影": ["维克黑沙滩", "钻石沙滩", "雷尼斯岩柱", "教会山", "极光观测点"],
            "亲子": ["蓝湖温泉", "熔岩中心", "鲸鱼博物馆", "雷克雅未克港口"],
            "蜜月": ["蓝湖私享温泉", "南岸瀑布轻徒步", "景观餐厅", "精品冰川酒店"],
            "商务考察": ["雷克雅未克酒店考察", "地接社拜访", "冰川活动供应商会议", "可持续旅游项目交流"],
        },
    },
    "芬兰": {
        "route": "赫尔辛基、波尔沃、罗瓦涅米、拉普兰",
        "entry": "赫尔辛基",
        "regions": ["赫尔辛基", "波尔沃", "罗瓦涅米", "拉普兰", "伊纳里"],
        "spots": {
            "自然风光": ["芬兰堡", "努克西奥国家公园", "拉普兰森林", "伊纳里湖"],
            "摄影": ["赫尔辛基设计区", "波尔沃老城", "北极圈地标", "拉普兰雪原"],
            "亲子": ["圣诞老人村", "驯鹿农场", "哈士奇雪橇营地", "芬兰堡亲子步道"],
            "蜜月": ["玻璃穹顶酒店", "湖畔桑拿", "雪屋晚餐", "极光私享观测"],
            "商务考察": ["赫尔辛基会展中心", "设计品牌访谈", "酒店集团考察", "北欧旅游局交流"],
        },
    },
    "挪威": {
        "route": "奥斯陆、卑尔根、弗洛姆、松恩峡湾",
        "entry": "奥斯陆",
        "regions": ["奥斯陆", "卑尔根", "弗洛姆", "松恩峡湾", "罗弗敦群岛", "特罗姆瑟"],
        "spots": {
            "自然风光": ["松恩峡湾", "盖朗厄尔峡湾", "弗洛姆高山铁路", "哈当厄尔峡湾"],
            "摄影": ["布吕根码头", "罗弗敦渔村", "峡湾观景台", "大西洋之路"],
            "亲子": ["奥斯陆民俗博物馆", "峡湾游船", "弗洛姆铁路", "水族馆"],
            "蜜月": ["峡湾景观酒店", "卑尔根海港晚餐", "罗弗敦海景木屋"],
            "商务考察": ["奥斯陆旅游机构拜访", "峡湾游船供应商考察", "酒店资源踩线", "邮轮产品交流"],
        },
    },
    "瑞典": {
        "route": "斯德哥尔摩、乌普萨拉、基律纳、阿比斯库",
        "entry": "斯德哥尔摩",
        "regions": ["斯德哥尔摩", "乌普萨拉", "基律纳", "阿比斯库"],
        "spots": {
            "自然风光": ["斯德哥尔摩群岛", "阿比斯库国家公园", "基律纳雪原"],
            "摄影": ["老城 Gamla Stan", "地铁艺术站", "群岛日落", "阿比斯库极光点"],
            "亲子": ["瓦萨沉船博物馆", "斯堪森露天博物馆", "群岛短线游船"],
            "蜜月": ["群岛精品酒店", "老城慢游", "冰酒店体验"],
            "商务考察": ["会奖资源考察", "设计品牌访问", "目的地管理公司会谈"],
        },
    },
}

CUSTOMER_NAMES = ["林嘉宁", "周亦辰", "王思岚", "赵明煦", "何沐阳", "许安琪", "沈知远", "陆景澄"]
DEPARTURE_CITIES = ["上海", "北京", "广州", "深圳", "杭州", "成都"]
PREFERENCES = ["自然风光", "摄影", "亲子", "蜜月", "商务考察"]
SPECIALS = [
    "希望控制每日车程，安排中文司导。",
    "需提供可退改酒店方案，并准备英文版签证辅助材料。",
    "一位客人不吃牛羊肉，餐食需提前备注。",
    "同行有儿童，需减少高强度徒步项目。",
    "需保留酒店和地接供应商考察时间。",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_label_table(doc: Document, rows: list[tuple[str, str]], widths=(1.55, 4.75)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        set_cell_shading(cells[0], "F2F4F7")
        set_cell_text(cells[0], label, True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.color.rgb = RGBColor(31, 95, 115)


def budget_tier(budget: int) -> str:
    if budget >= 60000:
        return "高端型"
    if budget >= 30000:
        return "舒适型"
    return "经济型"


def hotel_desc(tier: str) -> str:
    return {
        "经济型": "三星至四星基础酒店或高评分公寓式住宿，优先控制总预算与交通成本。",
        "舒适型": "四星精品酒店或交通便利型商务酒店，兼顾早餐、位置和取消政策。",
        "高端型": "五星酒店、景观酒店或高端精品住宿，优先安排核心区域与弹性退改条款。",
    }[tier]


def make_day_plan(day: int, total_days: int, profile: dict, preference: str, customer: dict, tier: str) -> tuple[str, str]:
    regions = profile["regions"]
    spots = profile["spots"][preference]
    region = regions[(day - 1) % len(regions)]
    primary = spots[(day - 1) % len(spots)]
    secondary = spots[day % len(spots)]
    current_date = customer["date"] + timedelta(days=day - 1)
    date_text = current_date.strftime("%Y年%m月%d日")
    pace = {
        "自然风光": "以自然景观与轻户外体验为主。",
        "摄影": "优先安排清晨或傍晚光线窗口。",
        "亲子": "控制步行强度并预留亲子休息时段。",
        "蜜月": "安排慢节奏体验、私密用餐和景观停留。",
        "商务考察": "保留供应商会谈、酒店踩线和资料收集时间。",
    }[preference]
    if day == 1:
        return (
            f"Day{day}",
            f"{date_text}；{customer['departure']} - {profile['entry']}。抵达后接机并入住{tier}住宿，下午安排{region}适应性参访，重点覆盖{primary}。{pace}",
        )
    if day == total_days:
        return (
            f"Day{day}",
            f"{date_text}；{region} - {customer['departure']}。上午安排{primary}或{secondary}作为返程前补充参访，完成酒店退房、票据整理和签证辅助文书归档，随后送机返程。",
        )
    return (
        f"Day{day}",
        f"{date_text}；{profile['entry']}及{region}动态行程。上午参访{primary}，下午衔接{secondary}及周边体验，住宿按{tier}标准安排在{region}或交通便利区域。{pace}",
    )


def generate_customers() -> list[dict]:
    rng = random.SystemRandom()
    start = date(2026, 9, 10)
    destination_keys = list(DESTINATIONS.keys())
    customers = []
    used_names = rng.sample(CUSTOMER_NAMES, 5)
    for index, name in enumerate(used_names, start=1):
        destination = rng.choice(destination_keys)
        preference = rng.choice(PREFERENCES)
        days = rng.choice([7, 8, 9, 10, 12])
        budget = rng.choice([22000, 28000, 36000, 48000, 62000, 78000])
        travelers = rng.choice([2, 3, 4, 5, 6])
        customers.append(
            {
                "index": index,
                "name": name,
                "travelers": travelers,
                "departure": rng.choice(DEPARTURE_CITIES),
                "destination": destination,
                "route": DESTINATIONS[destination]["route"],
                "date": start + timedelta(days=rng.randint(0, 80)),
                "days": days,
                "budget": budget,
                "preference": preference,
                "special": rng.choice(SPECIALS),
            }
        )
    return customers


def add_customer_section(doc: Document, customer: dict) -> None:
    profile = DESTINATIONS[customer["destination"]]
    tier = budget_tier(customer["budget"])
    end_date = customer["date"] + timedelta(days=customer["days"] - 1)
    base_total = customer["budget"] * customer["travelers"]
    service_fee = round(base_total * 0.08)
    doc_fee = 1200
    total = base_total + service_fee + doc_fee
    doc_no = f"NH-TEST-{date.today().strftime('%Y%m%d')}-{customer['index']:02d}"

    add_heading(doc, f"客户 {customer['index']}：{customer['name']} 签证辅助文书", 1)
    intro = doc.add_paragraph()
    intro.add_run(f"文档编号：{doc_no}    生成日期：{date.today().strftime('%Y年%m月%d日')}").bold = True
    doc.add_paragraph(f"联系方式：{AGENCY['dept']} / {AGENCY['phone']} / {AGENCY['email']}")

    add_heading(doc, "1. 客户需求概览", 2)
    add_label_table(
        doc,
        [
            ("客户姓名", customer["name"]),
            ("出行人数", f"{customer['travelers']}人"),
            ("出发城市", customer["departure"]),
            ("目的地", f"{customer['destination']}：{customer['route']}"),
            ("出行日期", f"{customer['date'].strftime('%Y年%m月%d日')} 至 {end_date.strftime('%Y年%m月%d日')}"),
            ("行程天数", f"{customer['days']}天"),
            ("人均预算", f"¥{customer['budget']:,}"),
            ("住宿标准", f"{tier}：{hotel_desc(tier)}"),
            ("出行偏好", customer["preference"]),
            ("特殊需求", customer["special"]),
        ],
    )

    add_heading(doc, "2. 行程安排", 2)
    itinerary = [
        make_day_plan(day, customer["days"], profile, customer["preference"], customer, tier)
        for day in range(1, customer["days"] + 1)
    ]
    add_label_table(doc, itinerary, widths=(0.85, 5.45))
    doc.add_paragraph(f"说明：本次行程严格按 {customer['days']} 天动态生成，输出范围为 Day1-Day{customer['days']}。")

    add_heading(doc, "3. 费用说明", 2)
    add_label_table(
        doc,
        [
            ("预算分级", f"{tier}住宿与服务标准"),
            ("基础旅行服务费", f"¥{customer['budget']:,} / 人 x {customer['travelers']}人 = ¥{base_total:,}"),
            ("定制与文书服务费", f"¥{service_fee:,}"),
            ("材料处理费", f"¥{doc_fee:,}"),
            ("预估总价", f"¥{total:,}"),
        ],
    )

    add_heading(doc, "4. 签证材料清单", 2)
    for item in [
        "护照原件、身份证复印件、户口本复印件、婚姻状况证明。",
        "近6个月白底彩色证件照，规格按申根签证中心当前要求执行。",
        "在职证明英文版、营业执照副本复印件盖章或收入来源说明。",
        "近6个月银行流水、余额证明、房产或车辆等辅助资产材料。",
        "往返机票预订单、酒店确认函、境外保险单、英文行程单、费用说明及邀请函。",
        f"补充说明：{customer['special']}",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "5. 酒店确认函", 2)
    add_heading(doc, "HOTEL RESERVATION CONFIRMATION LETTER", 3)
    doc.add_paragraph("To: Visa Officer / Consular Section")
    doc.add_paragraph(f"Subject: Accommodation Confirmation for {customer['name']} and Travel Party")
    doc.add_paragraph(
        f"{AGENCY['name']} hereby confirms that accommodation arrangements are being coordinated for "
        f"{customer['name']} and accompanying travel party of {customer['travelers']} person(s), in connection "
        f"with their planned visit to {customer['destination']}."
    )
    doc.add_paragraph(
        f"The scheduled accommodation period is from {customer['date'].strftime('%Y-%m-%d')} to "
        f"{end_date.strftime('%Y-%m-%d')}. The proposed accommodation tier is {tier}, selected according "
        f"to the declared per-person budget and travel requirements."
    )
    add_label_table(
        doc,
        [
            ("Lead Guest", customer["name"]),
            ("Number of Guests", f"{customer['travelers']} person(s)"),
            ("Travel Period", f"{customer['date'].strftime('%Y-%m-%d')} - {end_date.strftime('%Y-%m-%d')}"),
            ("Accommodation Standard", f"{tier} / {hotel_desc(tier)}"),
            ("Reservation Status", "Pre-confirmed for visa support; final booking number to be issued by supplier."),
        ],
    )

    add_heading(doc, "6. 商务邀请函", 2)
    add_heading(doc, "BUSINESS INVITATION LETTER", 3)
    doc.add_paragraph("To: Visa Officer / Consular Section")
    doc.add_paragraph(f"Subject: Invitation for Business Travel and Tourism Service Inspection in {customer['destination']}")
    doc.add_paragraph(
        f"{AGENCY['name']} formally invites {customer['name']} and accompanying travel party of "
        f"{customer['travelers']} person(s) to visit {customer['destination']} during the period from "
        f"{customer['date'].strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}."
    )
    doc.add_paragraph(
        f"The visit purpose is matched to the declared preference profile: {customer['preference']}. "
        "The itinerary includes destination resource review, hotel and local service inspection, and selected travel product experience."
    )
    doc.add_paragraph(
        f"Issued by: {AGENCY['name']} / {AGENCY['cn_name']}\n"
        f"Contact: {AGENCY['phone']} / {AGENCY['email']}\n"
        "Remark: This invitation is provided for visa support and business travel documentation purposes."
    )


def build_docx(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI北欧定制游顾问 - 随机客户批量测试文书")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 32, 42)
    doc.add_paragraph(f"签发机构：{AGENCY['name']} / {AGENCY['cn_name']}")
    doc.add_paragraph(f"生成日期：{date.today().strftime('%Y年%m月%d日')}    测试样本：5位随机客户")
    doc.add_paragraph(
        "备注：本文件用于测试当前项目的动态文书生成能力。所有客户姓名、预算、偏好和需求均为随机生成的演示数据。"
    )

    customers = generate_customers()
    for index, customer in enumerate(customers):
        if index == 0:
            doc.add_paragraph()
        else:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_customer_section(doc, customer)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    build_docx(Path("output/AI北欧定制游顾问_5位随机客户测试文书.docx"))
